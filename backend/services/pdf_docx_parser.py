import io
import re

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes safely."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text_parts.append(extracted)
        return "\n".join(text_parts)
    except Exception as e:
        # Fallback regex string cleanup if needed
        return f"PDF Extraction Error: {str(e)}"

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes safely."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        return f"DOCX Extraction Error: {str(e)}"

def clean_text(raw_text: str) -> str:
    """Normalize whitespace and remove unprintable characters."""
    if not raw_text:
        return ""
    text = re.sub(r'[\r\t]', ' ', raw_text)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()
